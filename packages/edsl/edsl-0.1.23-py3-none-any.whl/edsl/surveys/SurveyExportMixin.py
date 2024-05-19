"""A mixin class for exporting surveys to different formats."""
from docx import Document
from typing import Union
import black


class SurveyExportMixin:
    """A mixin class for exporting surveys to different formats."""

    def docx(self, filename=None) -> Union["Document", None]:
        """Generate a docx document for the survey."""
        doc = Document()
        doc.add_heading("EDSL Survey")
        doc.add_paragraph(f"\n")
        for index, question in enumerate(self._questions):
            h = doc.add_paragraph()  # Add question as a paragraph
            h.add_run(f"Question {index + 1} ({question.question_name})").bold = True
            h.add_run(f"; {question.question_type}").italic = True
            p = doc.add_paragraph()
            p.add_run(question.question_text)
            if question.question_type == "linear_scale":
                for key, value in getattr(question, "option_labels", {}).items():
                    doc.add_paragraph(str(key) + ": " + str(value), style="ListBullet")
            else:
                if hasattr(question, "question_options"):
                    for option in getattr(question, "question_options", []):
                        doc.add_paragraph(str(option), style="ListBullet")
        if filename:
            doc.save(filename)
            print("The survey has been saved to", filename)
            return
        return doc

    def code(self, filename: str = None, survey_var_name: str = "survey") -> list[str]:
        """Create the Python code representation of a survey.

        :param filename: The name of the file to save the code to.
        :param survey_var_name: The name of the survey variable.
        """
        header_lines = ["from edsl.surveys.Survey import Survey"]
        header_lines.append("from edsl import Question")
        lines = ["\n".join(header_lines)]
        for question in self._questions:
            question.question_text = question["question_text"].replace("\n", " ")
            # remove dublicate spaces
            question.question_text = " ".join(question.question_text.split())
            lines.append(f"{question.question_name} = " + repr(question))
        lines.append(
            f"{survey_var_name} = Survey(questions = [{', '.join(self.question_names)}])"
        )
        # return lines
        code_string = "\n".join(lines)
        formatted_code = black.format_str(code_string, mode=black.FileMode())

        if filename:
            print("The code has been saved to", filename)
            print("The survey itself is saved to 'survey' object")
            with open(filename, "w") as file:
                file.write(formatted_code)
            return

        return formatted_code

    def html(self, filename=None) -> str:
        """Generate the html for the survey."""
        html_text = []
        for question in self._questions:
            html_text.append(
                f"<p><b>{question.question_name}</b> ({question.question_type}): {question.question_text}</p>"
            )
            html_text.append("<ul>")
            for option in getattr(question, "question_options", []):
                html_text.append(f"<li>{option}</li>")
            html_text.append("</ul>")
        lines = "\n".join(html_text)
        if filename:
            print("The survey has been saved to", filename)
            with open(filename, "w") as file:
                file.write(lines)
            return
        return lines
